#!/usr/bin/env python3
"""
Debug script for document upload functionality
"""

import requests
import json
import os
import uuid
from dotenv import load_dotenv
from reportlab.pdfgen import canvas
from docx import Document as DocxDocument
import openpyxl
import csv
from pptx import Presentation

# Load environment variables
load_dotenv('/app/frontend/.env')
load_dotenv('/app/backend/.env')

# Backend URL configuration
BACKEND_URL = os.environ.get('REACT_APP_BACKEND_URL')
INTERNAL_API_URL = "http://localhost:8001/api"
API_URL = INTERNAL_API_URL

print(f"Testing document upload functionality at: {API_URL}")

def create_test_session():
    """Create a test session"""
    url = f"{API_URL}/sessions"
    payload = {"title": "Document Upload Debug Session"}
    
    response = requests.post(url, json=payload)
    if response.status_code != 200:
        print(f"Failed to create session: {response.status_code} - {response.text}")
        return None
    
    data = response.json()
    session_id = data["id"]
    print(f"✅ Created test session: {session_id}")
    return session_id

def test_supported_formats():
    """Test the supported formats endpoint"""
    print("\n=== Testing Supported Formats Endpoint ===")
    
    url = f"{API_URL}/supported-formats"
    response = requests.get(url)
    
    print(f"Status: {response.status_code}")
    if response.status_code == 200:
        data = response.json()
        print(f"Response: {json.dumps(data, indent=2)}")
        return data.get("supported_formats", [])
    else:
        print(f"Error: {response.text}")
        return []

def create_sample_files():
    """Create sample files for testing"""
    files = {}
    
    # Create PDF
    pdf_path = f"/app/debug_test_{uuid.uuid4().hex[:8]}.pdf"
    c = canvas.Canvas(pdf_path)
    c.drawString(100, 750, "Debug Test PDF Document")
    c.drawString(100, 700, "This is a test PDF for debugging document upload.")
    c.save()
    files['pdf'] = pdf_path
    
    # Create DOCX
    docx_path = f"/app/debug_test_{uuid.uuid4().hex[:8]}.docx"
    doc = DocxDocument()
    doc.add_heading('Debug Test DOCX', 0)
    doc.add_paragraph('This is a test DOCX for debugging.')
    doc.save(docx_path)
    files['docx'] = docx_path
    
    # Create TXT
    txt_path = f"/app/debug_test_{uuid.uuid4().hex[:8]}.txt"
    with open(txt_path, 'w') as f:
        f.write("Debug Test TXT Document\nThis is a test TXT for debugging.")
    files['txt'] = txt_path
    
    print(f"✅ Created sample files: {list(files.keys())}")
    return files

def test_document_upload(session_id, file_path, file_type):
    """Test uploading a document"""
    print(f"\n=== Testing {file_type.upper()} Upload ===")
    
    # Test new universal endpoint
    url = f"{API_URL}/sessions/{session_id}/upload-document"
    
    with open(file_path, "rb") as file:
        files = {"file": (f"test.{file_type}", file, f"application/{file_type}")}
        response = requests.post(url, files=files)
    
    print(f"Universal endpoint status: {response.status_code}")
    if response.status_code == 200:
        data = response.json()
        print(f"✅ Success: {json.dumps(data, indent=2)}")
        return True
    else:
        print(f"❌ Error: {response.text}")
        return False

def test_legacy_pdf_upload(session_id, pdf_path):
    """Test the legacy PDF upload endpoint"""
    print(f"\n=== Testing Legacy PDF Upload ===")
    
    url = f"{API_URL}/sessions/{session_id}/upload-pdf"
    
    with open(pdf_path, "rb") as file:
        files = {"file": ("test.pdf", file, "application/pdf")}
        response = requests.post(url, files=files)
    
    print(f"Legacy PDF endpoint status: {response.status_code}")
    if response.status_code == 200:
        data = response.json()
        print(f"✅ Success: {json.dumps(data, indent=2)}")
        return True
    else:
        print(f"❌ Error: {response.text}")
        return False

def test_unsupported_format(session_id):
    """Test uploading an unsupported format"""
    print(f"\n=== Testing Unsupported Format ===")
    
    # Create unsupported file
    unsupported_path = f"/app/debug_test_{uuid.uuid4().hex[:8]}.xyz"
    with open(unsupported_path, 'w') as f:
        f.write("This is an unsupported file format.")
    
    try:
        url = f"{API_URL}/sessions/{session_id}/upload-document"
        
        with open(unsupported_path, "rb") as file:
            files = {"file": ("test.xyz", file, "application/octet-stream")}
            response = requests.post(url, files=files)
        
        print(f"Unsupported format status: {response.status_code}")
        if response.status_code == 400:
            data = response.json()
            print(f"✅ Expected error: {json.dumps(data, indent=2)}")
            return True
        else:
            print(f"❌ Unexpected response: {response.text}")
            return False
    finally:
        if os.path.exists(unsupported_path):
            os.remove(unsupported_path)

def check_session_content(session_id):
    """Check if session was updated with document content"""
    print(f"\n=== Checking Session Content ===")
    
    url = f"{API_URL}/sessions"
    response = requests.get(url)
    
    if response.status_code != 200:
        print(f"❌ Failed to get sessions: {response.text}")
        return False
    
    sessions = response.json()
    test_session = next((s for s in sessions if s["id"] == session_id), None)
    
    if not test_session:
        print(f"❌ Test session not found")
        return False
    
    print(f"Session data: {json.dumps(test_session, indent=2)}")
    
    # Check for document fields
    has_document = any(key in test_session for key in ['document_content', 'pdf_content'])
    if has_document:
        print("✅ Session contains document content")
        return True
    else:
        print("❌ Session does not contain document content")
        return False

def main():
    """Main debug function"""
    print("="*80)
    print("DOCUMENT UPLOAD FUNCTIONALITY DEBUG")
    print("="*80)
    
    # Test supported formats
    supported_formats = test_supported_formats()
    
    # Create test session
    session_id = create_test_session()
    if not session_id:
        print("❌ Cannot proceed without session")
        return
    
    # Create sample files
    sample_files = create_sample_files()
    
    results = {}
    
    try:
        # Test each supported format
        for file_type in ['pdf', 'docx', 'txt']:
            if file_type in sample_files:
                success = test_document_upload(session_id, sample_files[file_type], file_type)
                results[f"upload_{file_type}"] = success
        
        # Test legacy PDF endpoint
        if 'pdf' in sample_files:
            success = test_legacy_pdf_upload(session_id, sample_files['pdf'])
            results["legacy_pdf"] = success
        
        # Test unsupported format
        success = test_unsupported_format(session_id)
        results["unsupported_format"] = success
        
        # Check session content
        success = check_session_content(session_id)
        results["session_content"] = success
        
    finally:
        # Clean up files
        for file_path in sample_files.values():
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except:
                pass
        
        # Clean up session
        try:
            requests.delete(f"{API_URL}/sessions/{session_id}")
        except:
            pass
    
    # Print summary
    print("\n" + "="*80)
    print("DEBUG SUMMARY")
    print("="*80)
    
    for test_name, success in results.items():
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{test_name}: {status}")
    
    total_tests = len(results)
    passed_tests = sum(1 for success in results.values() if success)
    success_rate = (passed_tests / total_tests) * 100 if total_tests > 0 else 0
    
    print(f"\nOverall Success Rate: {success_rate:.1f}% ({passed_tests}/{total_tests})")
    
    if success_rate >= 80:
        print("✅ DOCUMENT UPLOAD FUNCTIONALITY IS WORKING WELL")
    elif success_rate >= 60:
        print("⚠️  DOCUMENT UPLOAD FUNCTIONALITY HAS SOME ISSUES")
    else:
        print("❌ DOCUMENT UPLOAD FUNCTIONALITY NEEDS SIGNIFICANT FIXES")

if __name__ == "__main__":
    main()